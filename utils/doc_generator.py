from docx import Document

from datetime import datetime

from utils.logger import logger


def create_doc(title, content, filename):

    try:

        logger.info(f"Creating document: {filename}")

        doc = Document()

        doc.add_heading(title, level=1)

        doc.add_paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )

        doc.add_paragraph("")

        lines = content.split("\n")

        for line in lines:

            line = line.strip()

            if not line:
                continue

            if line.startswith("# "):

                doc.add_heading(
                    line.replace("# ", ""),
                    level=1
                )

            elif line.startswith("## "):

                doc.add_heading(
                    line.replace("## ", ""),
                    level=2
                )

            elif line.startswith("### "):

                doc.add_heading(
                    line.replace("### ", ""),
                    level=3
                )

            else:

                doc.add_paragraph(line)

        doc.save(filename)

        logger.info(f"Saved document: {filename}")

    except Exception as e:

        logger.error(f"DOCX Error: {str(e)}")