import os
import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
from serpapi import GoogleSearch
from docx import Document
from datetime import datetime


load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY")

# =========================
# Gemini CLIENT
# =========================
genai.configure(api_key=GEMINI_API_KEY)

model = genai.GenerativeModel("gemini-2.5-flash")

# =========================
# STREAMLIT CONFIG
# =========================
st.set_page_config(
    page_title="AI Teaching Agent Team",
    layout="wide"
)

st.title("👨‍🏫 AI Teaching Agent Team")
st.markdown("""
A collaborative AI teaching system where multiple AI agents work together like a professional teaching faculty.

### Agents:
-  Professor Agent
-  Academic Advisor Agent
-  Research Librarian Agent
-  Teaching Assistant Agent
""")


if "generated" not in st.session_state:
    st.session_state.generated = False


topic = st.text_input(
    "Enter a topic you want to learn",
    placeholder="Example: Machine Learning"
)

generate_btn = st.button(" Generate Learning Plan")



  
def ask_ai(system_prompt, user_prompt):

    full_prompt = f"""
    {system_prompt}

    User Request:
    {user_prompt}
    """

    response = model.generate_content(full_prompt)

    return response.text
   


def search_resources(topic):

    if not SERPAPI_API_KEY:
        return "SERPAPI KEY NOT FOUND"

    params = {
        "engine": "google",
        "q": f"best resources to learn {topic}",
        "api_key": SERPAPI_API_KEY,
        "num": 10
    }

    search = GoogleSearch(params)
    results = search.get_dict()

    resources = []

    if "organic_results" in results:

        for item in results["organic_results"][:10]:

            title = item.get("title", "No Title")
            link = item.get("link", "")
            snippet = item.get("snippet", "")

            resources.append(
                f"""
Title: {title}

Link: {link}

Description: {snippet}
"""
            )

    return "\n\n".join(resources)


def create_doc(title, content, filename):

    doc = Document()

    doc.add_heading(title, level=1)

    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    doc.add_paragraph("")

    lines = content.split("\n")

    for line in lines:

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)

        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)

        else:
            doc.add_paragraph(line)

    doc.save(filename)



if generate_btn and topic:

    with st.spinner(" AI Agents are collaborating..."):

       
        professor_prompt = f"""
You are Professor Agent.

Your job:
Create a COMPLETE knowledge base document for learning {topic}.

Requirements:
- Beginner friendly
- Structured properly
- Use headings and subheadings
- Include examples
- Include real world applications
- Include important terminology
- Include summary section
- Include advanced concepts

Format properly using markdown headings.
"""

        professor_output = ask_ai(
            professor_prompt,
            f"Create a complete knowledge base for {topic}"
        )

       

        advisor_prompt = f"""
You are Academic Advisor Agent.

Create a complete learning roadmap for {topic}.

Requirements:
- Beginner to advanced roadmap
- Weekly milestones
- Time estimates
- Prerequisites
- Project recommendations
- Interview preparation roadmap
- Revision strategy

Use proper markdown headings.
"""

        advisor_output = ask_ai(
            advisor_prompt,
            f"Create a learning roadmap for {topic}"
        )

       

        resources = search_resources(topic)

        librarian_prompt = f"""
You are Research Librarian Agent.

Below are collected web resources.

{resources}

Your task:
- Organize resources properly
- Categorize by difficulty
- Mention why resource is useful
- Include books, videos, documentation, tutorials
- Create a structured resource guide

Use markdown headings.
"""

        librarian_output = ask_ai(
            librarian_prompt,
            f"Create a resource guide for learning {topic}"
        )

       
        ta_prompt = f"""
You are Teaching Assistant Agent.

Create a practice workbook for {topic}.

Requirements:
- Beginner exercises
- Intermediate exercises
- Advanced exercises
- Coding challenges
- MCQs
- Assignments
- Mini projects
- Include answers and explanations

Use markdown headings.
"""

        ta_output = ask_ai(
            ta_prompt,
            f"Create a practice workbook for {topic}"
        )

       

        os.makedirs("generated_docs", exist_ok=True)

        professor_file = f"generated_docs/{topic}_knowledge_base.docx"
        advisor_file = f"generated_docs/{topic}_roadmap.docx"
        librarian_file = f"generated_docs/{topic}_resources.docx"
        ta_file = f"generated_docs/{topic}_practice_workbook.docx"

        create_doc(
            f"{topic} Knowledge Base",
            professor_output,
            professor_file
        )

        create_doc(
            f"{topic} Learning Roadmap",
            advisor_output,
            advisor_file
        )

        create_doc(
            f"{topic} Resource Guide",
            librarian_output,
            librarian_file
        )

        create_doc(
            f"{topic} Practice Workbook",
            ta_output,
            ta_file
        )

        st.session_state.generated = True

      
        st.success("✅ AI Teaching Team completed the learning package!")

        tab1, tab2, tab3, tab4 = st.tabs([
            "🧠 Professor",
            "🗺️ Advisor",
            "📚 Librarian",
            "✍️ Teaching Assistant"
        ])

        with tab1:
            st.markdown(professor_output)

        with tab2:
            st.markdown(advisor_output)

        with tab3:
            st.markdown(librarian_output)

        with tab4:
            st.markdown(ta_output)

     
        st.divider()
        st.subheader("📥 Download Documents")

        col1, col2 = st.columns(2)

        with col1:

            with open(professor_file, "rb") as f:
                st.download_button(
                    "Download Knowledge Base",
                    data=f,
                    file_name=os.path.basename(professor_file)
                )

            with open(librarian_file, "rb") as f:
                st.download_button(
                    "Download Resource Guide",
                    data=f,
                    file_name=os.path.basename(librarian_file)
                )

        with col2:

            with open(advisor_file, "rb") as f:
                st.download_button(
                    "Download Learning Roadmap",
                    data=f,
                    file_name=os.path.basename(advisor_file)
                )

            with open(ta_file, "rb") as f:
                st.download_button(
                    "Download Practice Workbook",
                    data=f,
                    file_name=os.path.basename(ta_file)
                )


with st.sidebar:

    st.header("⚙️ Configuration")

    st.markdown("""
### Required Environment Variables

Create a `.env` file:

```env
OPENAI_API_KEY=your_openai_key
SERPAPI_API_KEY=your_serpapi_key""")