import os
import pickle
from pathlib import Path

from dotenv import load_dotenv

from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

import google.generativeai as genai

from rank_bm25 import BM25Okapi

import cohere


try:
    from src.config.prompts import PROMPTS

except ImportError:
    from config.prompts import PROMPTS


ROOT_DIR = Path(__file__).resolve().parent.parent

GENERATED_DOCS_DIR = ROOT_DIR / "generated_docs"

VECTOR_DB_DIR = ROOT_DIR / "vector_db"


# LOAD ENV
load_dotenv(ROOT_DIR / ".env")

api_key = os.getenv("GEMINI_API_KEY")

cohere_key = os.getenv("CO_API_KEY")


# GEMINI
genai.configure(api_key=api_key)

co = cohere.Client(cohere_key)


# EMBEDDINGS
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)


# LOAD DOCX
def load_docx(path):

    doc = DocxDocument(path)

    text = "\n".join(
        para.text
        for para in doc.paragraphs
    )

    return Document(
        page_content=text,
        metadata={
            "source": str(path)
        }
    )


# LOAD ALL DOCS FOR A TOPIC
def load_topic_documents(topic_name):

    topic_path = os.path.join(
        GENERATED_DOCS_DIR,
        topic_name
    )

    documents = []

    if not os.path.exists(topic_path):

        return documents

    for file in os.listdir(topic_path):

        if file.endswith(".docx"):

            full_path = os.path.join(
                topic_path,
                file
            )

            try:

                doc = load_docx(full_path)

                documents.append(doc)

            except Exception as e:

                print(f"Error loading {file}: {e}")

    return documents


# HIERARCHICAL RETRIEVER
class HierarchicalRetriever:

    def __init__(self, embeddings, persist_dir):

        self.embeddings = embeddings

        self.persist_dir = persist_dir

        self.child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=300,
            chunk_overlap=50
        )

        self.parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=200
        )

        self.vectorstore = None

        self.docstore = {}

        self.docstore_path = os.path.join(
            persist_dir,
            "parent_docstore.pkl"
        )

    def add_documents(self, documents):

        parent_docs = self.parent_splitter.split_documents(
            documents
        )

        all_children = []

        for i, parent in enumerate(parent_docs):

            parent_id = f"parent_{i}"

            self.docstore[parent_id] = {
                "page_content": parent.page_content,
                "metadata": parent.metadata,
            }

            sub_children = self.child_splitter.split_documents(
                [parent]
            )

            for child in sub_children:

                child.metadata["parent_id"] = parent_id

                all_children.append(child)

        self.vectorstore = Chroma.from_documents(
            all_children,
            embedding=self.embeddings,
            persist_directory=self.persist_dir,
            collection_name="hierarchical_children",
        )

        self._save_docstore()

        print(
            f"Indexed {len(parent_docs)} parents and {len(all_children)} children."
        )

    def _save_docstore(self):

        with open(self.docstore_path, "wb") as f:

            pickle.dump(self.docstore, f)

    def _load_docstore(self):

        if os.path.exists(self.docstore_path):

            with open(self.docstore_path, "rb") as f:

                self.docstore = pickle.load(f)

            print("Loaded parent context from disk.")

    def get_relevant_documents(self, query, k=6):

        if self.vectorstore is None:

            return []

        results = self.vectorstore.similarity_search(
            query,
            k=k * 3
        )

        parents = []

        seen_ids = set()

        for child in results:

            pid = child.metadata.get("parent_id")

            if pid and pid not in seen_ids:

                item = self.docstore.get(pid)

                if item:

                    parents.append(
                        Document(
                            page_content=item["page_content"],
                            metadata=item["metadata"],
                        )
                    )

                    seen_ids.add(pid)

        return parents[:k]


# GET / CREATE RETRIEVER
def get_hierarchical_components(topic_name, documents):

    persist_dir = os.path.join(
        VECTOR_DB_DIR,
        topic_name
    )

    retriever = HierarchicalRetriever(
        embeddings,
        persist_dir
    )

    if os.path.exists(persist_dir):

        print(f"Loading existing vector DB for {topic_name}")

        retriever.vectorstore = Chroma(
            persist_directory=persist_dir,
            embedding_function=embeddings,
            collection_name="hierarchical_children",
        )

        retriever._load_docstore()

    else:

        print(f"Creating vector DB for {topic_name}")

        os.makedirs(
            persist_dir,
            exist_ok=True
        )

        retriever.add_documents(documents)

    return retriever


# RERANK
def rerank(query, docs, top_n=3):

    if not docs:

        return []

    texts = [
        doc.page_content
        for doc in docs
    ]

    results = co.rerank(
        model="rerank-english-v3.0",
        query=query,
        documents=texts,
        top_n=top_n
    )

    return [
        docs[res.index]
        for res in results.results
        if res.relevance_score > 0.20
    ]


# QUERY EXPANSION
def expand_query(query):

    prompt = PROMPTS["query_expansion"].format(
        question=query
    )

    model = genai.GenerativeModel(
        "gemini-2.5-flash"
    )

    return model.generate_content(
        prompt
    ).text.strip()


# MAIN QA FUNCTION
def ask_question(question, topic_name):

    documents = load_topic_documents(
        topic_name
    )

    if not documents:

        return "No documents found for this topic.", []

    retriever = get_hierarchical_components(
        topic_name,
        documents
    )

    child_splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=50
    )

    child_chunks = child_splitter.split_documents(
        documents
    )

    bm25 = BM25Okapi(
        [doc.page_content.split() for doc in child_chunks]
    )

    expanded_query = expand_query(question)

    # VECTOR SEARCH
    vector_parents = retriever.get_relevant_documents(
        expanded_query
    )

    # BM25 SEARCH
    tokenized_query = expanded_query.split()

    bm25_indices = sorted(
        range(len(bm25.get_scores(tokenized_query))),
        key=lambda i: bm25.get_scores(tokenized_query)[i],
        reverse=True,
    )[:5]

    bm25_children = [
        child_chunks[i]
        for i in bm25_indices
    ]

    # HYBRID MERGE
    combined = vector_parents + bm25_children

    unique_docs = list({
        doc.page_content: doc
        for doc in combined
    }.values())

    # RERANK
    docs = rerank(
        question,
        unique_docs
    )

    if not docs:

        return "The documents do not contain this information.", []

    # BUILD CONTEXT
    context = ""

    for i, doc in enumerate(docs):

        source = os.path.basename(
            doc.metadata.get("source", "Unknown")
        )

        context += (
            f"\n--- Source {i+1}: {source} ---\n"
        )

        context += f"{doc.page_content}\n\n"

    # FINAL PROMPT
    prompt = PROMPTS["rag_answer"].format(
        context=context,
        question=question
    )

    model = genai.GenerativeModel(
        "gemini-2.5-flash"
    )

    response = model.generate_content(
        prompt
    )

    answer_text = response.text

    return answer_text, docs


# TERMINAL MODE
if __name__ == "__main__":

    print(
        "Teaching Agent RAG System Ready!"
    )

    while True:

        topic = input(
            "\nTopic name: "
        )

        query = input(
            "Ask a question: "
        )

        if query.lower() == "exit":

            break

        answer, _ = ask_question(
            query,
            topic
        )

        print(
            "\n" + "=" * 80 +
            "\nANSWER:\n" +
            answer +
            "\n" + "=" * 80
        )